from docx import Document

document = Document()

# Them thong tin
document.add_heading("MY DOC", level=0)
document.add_heading("Chương 1", level=1)

# Doan
document.add_paragraph("It was a dark and stormy night.")

# Bang
records = (
    (3, '101', 'Spam'),
    (7, '422', 'Eggs'),
    (4, '631', 'Spam, spam, eggs, and spam')
)

table = document.add_table(rows=1, cols=3)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Qty'
hdr_cells[1].text = 'Id'
hdr_cells[2].text = 'Desc'
for qty, id, desc in records:
    row_cells = table.add_row().cells
    row_cells[0].text = str(qty)
    row_cells[1].text = id
# table.style = 'LightShading-Accent1'


# Hinh

# Save file
document.save("Python.docx")